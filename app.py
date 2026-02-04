from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import RGBColor
import io
import os
import re

app = Flask(__name__)

@app.route('/')
def home():
    return jsonify({
        'status': 'online',
        'message': 'CBT DOCX Processor API v2',
        'endpoints': {
            '/process-docx': 'POST - Process DOCX with replacements'
        }
    })

def replace_text_in_paragraph(paragraph, replacements_dict):
    """Replace text in paragraph, handling runs properly"""
    # Get full paragraph text
    full_text = paragraph.text
    
    # Check if replacement is needed
    needs_replacement = any(key in full_text for key in replacements_dict.keys())
    
    if not needs_replacement:
        return
    
    # Perform replacements on full text
    new_text = full_text
    for key, value in replacements_dict.items():
        new_text = new_text.replace(key, str(value))
    
    # Clear all runs
    for run in paragraph.runs:
        run.text = ''
    
    # Add new text to first run (preserves formatting)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        # If no runs exist, add text directly
        paragraph.add_run(new_text)

@app.route('/process-docx', methods=['POST'])
def process_docx():
    try:
        # Get file and replacements from request
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        docx_file = request.files['file']
        replacements = request.form.get('replacements')
        
        if not replacements:
            return jsonify({'error': 'No replacements provided'}), 400
        
        # Parse replacements JSON
        import json
        replacements_dict = json.loads(replacements)
        
        # Load DOCX
        doc = Document(docx_file)
        
        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements_dict)
        
        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements_dict)
        
        # Replace text in headers and footers
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                replace_text_in_paragraph(paragraph, replacements_dict)
            # Footer
            for paragraph in section.footer.paragraphs:
                replace_text_in_paragraph(paragraph, replacements_dict)
        
        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # Return processed file
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='processed.docx'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'healthy'}), 200

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))
    app.run(host='0.0.0.0', port=port)
